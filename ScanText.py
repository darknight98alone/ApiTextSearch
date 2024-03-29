import pytesseract
from handleTable import IOU,printImage
import cv2
import imutils
import numpy as np
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def checkIOUwithAboveRow(cell1,cell2):
    (x,y,w,h) = cell1
    (x1,y1,w1,h1) = cell2
    if (y1<=y and y<=y1+h1-5) or (y<=y1 and y+h-5>=y1):
        return True
    return False

def checkIOUwithAboveCell(cell1,cell2):
    (x,y,w,h) = cell1
    (x1,y1,w1,h1) = cell2
    if (x<=x1 and x1<=x+w) or (x1<=x and x <= x1+w1):
        return True
    return False

def layoutDocument(image,document):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
    # assign a rectangle kernel size
    kernel = np.ones((3, 3), 'uint8')
    par_img = cv2.dilate(thresh, kernel, iterations=6)
    if imutils.is_cv2() or imutils.is_cv4():
        (conts, hierarchy) = cv2.findContours(par_img.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    elif imutils.is_cv3():
        (_, conts, hierarchy) = cv2.findContours(par_img.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    listResult = []
    if len(conts) > 0:
        conts = imutils.contours.sort_contours(conts)[0]
        for i in range(len(conts)):
            (x, y, w, h) = cv2.boundingRect(conts[i])
            skip = False
            for temp in listResult:
                for box in temp:
                    if IOU(box,(x,y,w,h)):
                        skip = True
                        break
            if skip == False:
                over = False
                for index,temp in enumerate(listResult):
                    (x1,y1,w1,h1) = temp[0]
                    if abs(y1-y)<=5:
                        listResult[index].append((x,y,w,h))
                        over = True
                if over == False:
                    listResult.append([(x,y,w,h)])
        ## sort
    for index,_ in enumerate(listResult):
        listResult[index] = sorted(listResult[index], key=lambda x: x[0])
    listResult = sorted(listResult,key=lambda x: x[0][1])
    
    for temp in listResult:
        if len(temp) == 1:
            (x,y,w,h) = temp[0]
            if w > image.shape[1]/3: ## kiem tra align
                align = WD_TABLE_ALIGNMENT.LEFT
                print("left")
            else:
                align = WD_TABLE_ALIGNMENT.CENTER
        else:
            align = WD_TABLE_ALIGNMENT.CENTER

        table = document.add_table(rows=1, cols=len(temp))
        skip = False
        pass1 = False
        for index,(x,y,w,h) in enumerate(temp):
            if skip == False:
                row_cells = table.rows[0].cells
                if index<len(temp)-1:
                    if pass1 == False and checkIOUwithAboveRow(listResult[index+1][0],(x,y,w,h)):
                        pass1 = True
                    if pass1 == True:
                        for val in listResult[index+1]:
                            if checkIOUwithAboveCell((x,y,w,h),val):
                                crop = image[y:y+h,x:x+w]
                                string = pytesseract.image_to_string(crop,lang='vie')+"\n"
                                cv2.rectangle(image, (x, y), (x + w, y + h), 255, 1)
                                (x,y,w,h) = val
                                crop = image[y:y+h,x:x+w]
                                string = string + pytesseract.image_to_string(crop,lang='vie')
                                cv2.rectangle(image, (x, y), (x + w, y + h), 255, 1)
                                pass1 = False
                                skip = True
                                break
                if skip == False:   
                    crop = image[y:y+h,x:x+w]
                    string = pytesseract.image_to_string(crop,lang='vie')
                    cv2.rectangle(image, (x, y), (x + w, y + h), 255, 1)
                p = row_cells[index].add_paragraph(string)
                p.alignment = align
                printImage(image)
                print(string)
            else:
                skip = False
    return listResult,document    

def GetTextLayout(listResult,listBigBox,img,docName):
    ## open or create docx file
    if os.path.exists(docName) == False:
        dc = Document()
        dc.save(docName)
    document = Document(docName)

    result = []
    (height,_) = img.shape[:2]
    if len(listBigBox)==0:
        layoutDocument(img,document)
        document.save(docName)
        result.append(pytesseract.image_to_string(img,lang='vie')+"\n")
        return result
    bigBoxTemp = []
    listYCoord = []
    listYCoord.append(0)
    for (_,y,_,h) in listBigBox:
        bigBoxTemp.append((y,y+h))
        listYCoord.append(y)
        listYCoord.append(y+h)
    listYCoord.append(height)
    for index,y1 in enumerate(listYCoord):
        if index == len(listYCoord)-1:
            break
        y2 = listYCoord[index+1]
        if (y1,y2) not in bigBoxTemp:
            crop = img[y1:y2,:] # not table
            idx = (crop.flatten()<5)
            temp = sum(idx[:])
            if temp <2:
                continue
            layoutDocument(crop,document)
            result.append(pytesseract.image_to_string(crop,lang='vie')+"\n")
        else:
            index = 0
            box = listBigBox.pop(0)
            for temp in listResult:
                if IOU(temp[0],box):
                    index = index +1
                    table = document.add_table(rows=1, cols=len(listResult[0]))
                    for i,(x,y,w,h) in enumerate(temp):
                        crop = img[y:y+h,x:x+w]
                        size = int(crop.shape[0]*1.5)
                        if size < 100:
                            size = 100
                        crop = imutils.resize(crop,height=size) ## 
                        crop = cv2.cvtColor(crop, cv2.COLOR_BGR2GRAY)
                        idx = (crop.flatten()<5)
                        temp = sum(idx[:])
                        if temp <2:
                            continue
                        string  = pytesseract.image_to_string(crop,lang='vie')
                        if len(string) == 0:
                            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3,3))
                            crop = cv2.erode(crop, kernel, iterations=1)
                            while size<=200:
                                string  = pytesseract.image_to_string(crop,lang='eng',config='--psm 10')
                                if ("l" in string or "I" in string ) and len(string)<3:
                                    string = "1"
                                if string.isdigit() or (len(string)>=2 and "," not in string and "'" not in string):
                                    break
                                size = size + 10
                                crop = imutils.resize(crop,height=size)
                        string = string + " "
                        string = string.replace("\n"," ")
                        row_cells = table.rows[0].cells
                        p = row_cells[i].add_paragraph(string)
                        p.alignment = WD_TABLE_ALIGNMENT.CENTER
                        # result.append(string)
                    result.append("\n")
                else:
                    break
            listResult = listResult[index:]
    document.save(docName)
    return result

## stable get text    
def GetText(listResult,listBigBox,img):
    result = []
    (height,_) = img.shape[:2]
    if len(listBigBox)==0:
        result.append(pytesseract.image_to_string(img,lang='vie')+"\n")
        return result
    bigBoxTemp = []
    listYCoord = []
    listYCoord.append(0)
    for (_,y,_,h) in listBigBox:
        bigBoxTemp.append((y,y+h))
        listYCoord.append(y)
        listYCoord.append(y+h)
    listYCoord.append(height)
    for index,y1 in enumerate(listYCoord):
        if index == len(listYCoord)-1:
            break
        y2 = listYCoord[index+1]
        if (y1,y2) not in bigBoxTemp:
            crop = img[y1:y2,:] # not table
            idx = (crop.flatten()<5)
            temp = sum(idx[:])
            if temp <2:
                continue
            result.append(pytesseract.image_to_string(crop,lang='vie')+"\n")
        else:
            index = 0
            box = listBigBox.pop(0)
            for temp in listResult:
                if IOU(temp[0],box):
                    index = index +1
                    for (x,y,w,h) in temp:
                        crop = img[y:y+h,x:x+w]
                        size = int(crop.shape[0]*1.5)
                        if size < 100:
                            size = 100
                        crop = imutils.resize(crop,height=size) ## 
                        crop = cv2.cvtColor(crop, cv2.COLOR_BGR2GRAY)
                        idx = (crop.flatten()<5)
                        temp = sum(idx[:])
                        if temp <2:
                            continue
                        string  = pytesseract.image_to_string(crop,lang='vie')
                        if len(string) == 0:
                            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3,3))
                            crop = cv2.erode(crop, kernel, iterations=1)
                            while size<=200:
                                string  = pytesseract.image_to_string(crop,lang='eng',config='--psm 10')
                                if ("l" in string or "I" in string ) and len(string)<3:
                                    string = "1"
                                if string.isdigit() or (len(string)>=2 and "," not in string and "'" not in string):
                                    break
                                size = size + 10
                                crop = imutils.resize(crop,height=size)
                        string = string + " "
                        string = string.replace("\n"," ")
                        result.append(string)
                    result.append("\n")
                else:
                    break
            listResult = listResult[index:]
    return result